"""DOCX format converter for DocFlow."""

import re
from typing import List, Optional, Tuple

from ..types import ConvertResult, ExtractedImage


class DOCXConverter:
    """Converts between DOCX and Markdown formats.
    
    Requires python-docx: pip install python-docx
    
    Example:
        >>> converter = DOCXConverter()
        >>> result = converter.to_markdown(docx_bytes, "document.docx")
        >>> print(result.content)
    """

    def __init__(
        self,
        extract_images: bool = True,
        preserve_tables: bool = True,
    ) -> None:
        """Initialize DOCX converter.
        
        Args:
            extract_images: Extract embedded images.
            preserve_tables: Preserve table structure.
        """
        self.extract_images = extract_images
        self.preserve_tables = preserve_tables

    def to_markdown(
        self,
        docx_data: bytes,
        filename: str = "document.docx",
        include_metadata: bool = True,
    ) -> ConvertResult:
        """Convert DOCX to Markdown.
        
        Args:
            docx_data: DOCX file bytes.
            filename: Original filename.
            include_metadata: Include YAML frontmatter.
            
        Returns:
            ConvertResult with markdown and extracted images.
        """
        try:
            from docx import Document
            from docx.table import Table
            from docx.text.paragraph import Paragraph
            from io import BytesIO
        except ImportError:
            return ConvertResult(
                success=False,
                error="python-docx is required: pip install python-docx"
            )

        try:
            doc = Document(BytesIO(docx_data))
            md_parts = []
            images: List[ExtractedImage] = []
            image_counter = 0

            # Metadata
            if include_metadata:
                md_parts.append("---")
                md_parts.append(f"source: {filename}")
                md_parts.append("format: docx")
                
                # Core properties
                if doc.core_properties.title:
                    md_parts.append(f"title: {doc.core_properties.title}")
                if doc.core_properties.author:
                    md_parts.append(f"author: {doc.core_properties.author}")
                if doc.core_properties.created:
                    md_parts.append(f"created: {doc.core_properties.created}")
                
                md_parts.append(f"paragraphs: {len(doc.paragraphs)}")
                md_parts.append(f"tables: {len(doc.tables)}")
                md_parts.append("---\n")

            # Process document body
            for element in doc.element.body:
                if element.tag.endswith('p'):  # Paragraph
                    para = Paragraph(element, doc)
                    md_parts.append(self._paragraph_to_md(para))
                elif element.tag.endswith('tbl'):  # Table
                    for table in doc.tables:
                        if table._tbl == element:
                            md_parts.append(self._table_to_md(table))
                            break

            # Extract images
            if self.extract_images:
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        try:
                            image_counter += 1
                            img_data = rel.target_part.blob
                            ext = rel.target_ref.split(".")[-1]
                            images.append(ExtractedImage(
                                data=img_data,
                                format=ext,
                                filename=f"image_{image_counter:03d}.{ext}",
                                caption=None,
                            ))
                        except:
                            pass

            content = "\n".join(md_parts)
            # Clean up excessive newlines
            content = re.sub(r'\n{3,}', '\n\n', content)

            return ConvertResult(
                success=True,
                content=content,
                format="docx",
                images=images,
                metadata={
                    "paragraphs": len(doc.paragraphs),
                    "tables": len(doc.tables),
                    "images": len(images),
                },
            )

        except Exception as e:
            return ConvertResult(success=False, error=str(e))

    def from_markdown(
        self,
        markdown: str,
        title: Optional[str] = None,
    ) -> ConvertResult:
        """Convert Markdown to DOCX.
        
        Args:
            markdown: Markdown content.
            title: Document title.
            
        Returns:
            ConvertResult with DOCX bytes.
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from io import BytesIO
        except ImportError:
            return ConvertResult(
                success=False,
                error="python-docx is required: pip install python-docx"
            )

        try:
            doc = Document()
            
            if title:
                doc.core_properties.title = title

            lines = markdown.split("\n")
            i = 0
            
            # Skip frontmatter
            if lines and lines[0].strip() == "---":
                i = 1
                while i < len(lines) and lines[i].strip() != "---":
                    i += 1
                i += 1

            in_code_block = False
            code_content = []
            table_lines = []
            in_table = False

            while i < len(lines):
                line = lines[i]
                
                # Code blocks
                if line.strip().startswith("```"):
                    if in_code_block:
                        para = doc.add_paragraph()
                        run = para.add_run("\n".join(code_content))
                        run.font.name = "Courier New"
                        run.font.size = Pt(10)
                        code_content = []
                    in_code_block = not in_code_block
                    i += 1
                    continue

                if in_code_block:
                    code_content.append(line)
                    i += 1
                    continue

                # Tables
                if line.strip().startswith("|") and line.strip().endswith("|"):
                    in_table = True
                    table_lines.append(line)
                    i += 1
                    continue
                elif in_table:
                    self._add_table_to_doc(doc, table_lines)
                    table_lines = []
                    in_table = False

                # Headings
                if line.startswith("#"):
                    level = len(line.split()[0])
                    text = line.lstrip("#").strip()
                    doc.add_heading(text, level=min(level, 9))
                # Lists
                elif line.strip().startswith(("- ", "* ", "+ ")):
                    text = line.strip()[2:]
                    doc.add_paragraph(text, style="List Bullet")
                elif re.match(r'^\d+\.\s', line.strip()):
                    text = re.sub(r'^\d+\.\s', '', line.strip())
                    doc.add_paragraph(text, style="List Number")
                # Regular paragraph
                elif line.strip():
                    para = doc.add_paragraph()
                    self._add_formatted_text(para, line)
                
                i += 1

            # Handle remaining table
            if table_lines:
                self._add_table_to_doc(doc, table_lines)

            output = BytesIO()
            doc.save(output)
            output.seek(0)

            return ConvertResult(
                success=True,
                content=output.getvalue(),
                format="docx",
            )

        except Exception as e:
            return ConvertResult(success=False, error=str(e))

    def _paragraph_to_md(self, para) -> str:
        """Convert paragraph to markdown."""
        text = para.text.strip()
        if not text:
            return ""

        style_name = para.style.name.lower() if para.style else ""

        # Headings
        if "heading 1" in style_name:
            return f"# {text}\n"
        elif "heading 2" in style_name:
            return f"## {text}\n"
        elif "heading 3" in style_name:
            return f"### {text}\n"
        elif "heading" in style_name:
            return f"#### {text}\n"
        elif "title" in style_name:
            return f"# {text}\n"
        # Lists
        elif "list" in style_name:
            return f"- {text}"
        # Regular text
        else:
            # Apply inline formatting
            result = text
            for run in para.runs:
                run_text = run.text
                if run.bold and run_text:
                    result = result.replace(run_text, f"**{run_text}**", 1)
                if run.italic and run_text:
                    result = result.replace(run_text, f"*{run_text}*", 1)
            return result + "\n"

    def _table_to_md(self, table) -> str:
        """Convert table to markdown."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append(cells)

        if not rows:
            return ""

        md_lines = []
        # Header
        md_lines.append("| " + " | ".join(rows[0]) + " |")
        md_lines.append("| " + " | ".join("---" for _ in rows[0]) + " |")
        # Data
        for row in rows[1:]:
            md_lines.append("| " + " | ".join(row) + " |")

        return "\n" + "\n".join(md_lines) + "\n"

    def _add_table_to_doc(self, doc, table_lines):
        """Add markdown table to DOCX document."""
        rows = []
        for line in table_lines:
            line = line.strip()
            if line.startswith("|") and line.endswith("|"):
                # Skip separator
                if set(line.replace("|", "").replace("-", "").replace(":", "").strip()) == set():
                    continue
                cells = [c.strip() for c in line.split("|")[1:-1]]
                rows.append(cells)

        if not rows:
            return

        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = "Table Grid"

        for i, row_data in enumerate(rows):
            for j, cell_data in enumerate(row_data):
                if j < len(table.rows[i].cells):
                    table.rows[i].cells[j].text = cell_data

    def _add_formatted_text(self, para, text):
        """Add formatted text to paragraph."""
        # Simple bold/italic handling
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = para.add_run(part[1:-1])
                run.italic = True
            else:
                para.add_run(part)
